"""Сборка отдельного .docx-файла с полным листингом исходного кода.

Получается документ в стилистике основного диплома (PDF):
- основной текст — Times New Roman 14 пт, интервал 1.5;
- листинги кода — Times New Roman 11 пт, интервал 1.0;
- заголовок «ПРИЛОЖЕНИЕ А. ЛИСТИНГИ ИСХОДНОГО КОДА» — прописными по центру;
- перед каждым листингом — подпись «Листинг — файл <path>. <Описание>».

Запуск: python docs/diploma/build_appendix_a.py
Результат: docs/diploma/appendix_a_full_listing.docx

Исходные коды берутся из каталога docs/diploma/source_listings/, чтобы
сборка не зависела от структуры рабочей копии репозитория.
"""
from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt


CODE_FONT = 'Times New Roman'
BODY_FONT = 'Times New Roman'


def _apply_font(run, name, size, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    # Подстраховка для кириллицы.
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), name)


def _build_document() -> Document:
    doc = Document()

    # Поля по ГОСТ: левое 30, правое 15, верх/низ 20 мм.
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(30)
        section.right_margin = Mm(15)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)

    # Базовый стиль документа.
    normal = doc.styles['Normal']
    normal.font.name = BODY_FONT
    normal.font.size = Pt(14)
    pf = normal.paragraph_format
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    return doc


def _add_heading(doc, text):
    """Заголовок приложения: прописные, по центру."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text.upper())
    _apply_font(run, BODY_FONT, 14, bold=True)


def _add_body(doc, text):
    """Обычный абзац основного текста."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    _apply_font(run, BODY_FONT, 14)


def _add_caption(doc, text):
    """Подпись перед листингом: без отступа, обычным шрифтом."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _apply_font(run, BODY_FONT, 14)


def _add_code(doc, text):
    """Листинг кода: TNR 11 пт, межстрочный 1.0, без абзацного отступа."""
    for line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else ' ')
        _apply_font(run, CODE_FONT, 11)


# Описание набора файлов, попадающих в приложение.
# Порядок и формулировки выдержаны в стиле основного документа.
FILES = [
    ('manage.py',
     'точка входа для управления Django-проектом'),
    ('config/settings.py',
     'настройки проекта (база данных, приложения, шаблоны, '
     'статика, медиа, AUTH_USER_MODEL)'),
    ('config/urls.py',
     'корневой файл маршрутизации, подключающий URL всех приложений'),

    ('users/models.py',
     'модель пользователя с указанием роли '
     '(администратор, тренер, клиент)'),
    ('users/signals.py',
     'сигнал автоматического создания профиля клиента или тренера '
     'при регистрации пользователя'),
    ('users/decorators.py',
     'декораторы проверки роли пользователя при доступе к разделам'),
    ('users/forms.py',
     'формы регистрации клиента и редактирования профиля'),
    ('users/views.py',
     'представления для регистрации, входа, дашбордов и профиля'),
    ('users/urls.py',
     'маршруты приложения users'),
    ('users/admin.py',
     'настройка отображения пользователей в админ-панели Django'),
    ('users/management/commands/seed_demo.py',
     'команда заполнения базы демонстрационными данными'),

    ('clients/models.py',
     'модель клиента (профиль, дата рождения, пол, контакты)'),
    ('clients/views.py',
     'представления списка клиентов и карточки клиента'),
    ('clients/urls.py',
     'маршруты приложения clients'),
    ('clients/admin.py',
     'настройка раздела «Клиенты» в админ-панели'),

    ('trainers/models.py',
     'модели тренера и специализации'),
    ('trainers/views.py',
     'представления списка тренеров и карточки тренера'),
    ('trainers/urls.py',
     'маршруты приложения trainers'),
    ('trainers/admin.py',
     'настройка раздела «Тренеры» в админ-панели'),

    ('subscriptions/models.py',
     'модели тарифного плана, абонемента и платежа'),
    ('subscriptions/views.py',
     'представления каталога тарифов и покупки абонемента'),
    ('subscriptions/urls.py',
     'маршруты приложения subscriptions'),
    ('subscriptions/admin.py',
     'настройка раздела «Абонементы» в админ-панели'),

    ('schedule/models.py',
     'модели залов, тренировок и записей клиентов на занятия'),
    ('schedule/views.py',
     'представления расписания, записи на тренировки и отметки посещений'),
    ('schedule/urls.py',
     'маршруты приложения schedule'),
    ('schedule/admin.py',
     'настройка раздела «Расписание» в админ-панели'),

    ('reports/forms.py',
     'формы выбора периода и фильтров для отчётов'),
    ('reports/services.py',
     'бизнес-логика построения отчётов '
     '(посещаемость, выручка, загрузка тренеров)'),
    ('reports/views.py',
     'представления для просмотра отчётов'),
    ('reports/urls.py',
     'маршруты приложения reports'),
]


def build(output: Path) -> None:
    doc = _build_document()
    _add_heading(doc, 'Приложение А. Листинги исходного кода')

    _add_body(doc,
              'В приложении приведён полный листинг исходного кода '
              'разработанной системы. Полный код проекта также размещён '
              'в репозитории на GitHub.')
    _add_body(doc,
              'Листинги выполнены шрифтом Times New Roman 11 пт с '
              'межстрочным интервалом 1,0. Файлы упорядочены по '
              'Django-приложениям проекта: общие настройки, '
              'пользователи, клиенты, тренеры, абонементы, '
              'расписание, отчёты. Файлы миграций, тестов, '
              'HTML-шаблоны и статические ресурсы в приложение не '
              'включены: миграции автоматически генерируются Django '
              'по описанию моделей, а шаблоны представляют собой '
              'разметку, а не исходный код программы.')

    listings_root = Path(__file__).parent / 'source_listings'

    for rel_path, description in FILES:
        path = listings_root / rel_path
        if not path.exists():
            continue
        source = path.read_text(encoding='utf-8').rstrip('\n')

        _add_caption(doc, f'Листинг — файл {rel_path}. {description}.')
        _add_code(doc, source)

    doc.save(output)
    print(f'Готово: {output}')


if __name__ == '__main__':
    root = Path(__file__).resolve().parent
    output = root / 'appendix_a_full_listing.docx'
    build(output)
