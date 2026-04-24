"""Сборщик пояснительной записки (диплом) в формат .docx.

Запуск: `python docs/diploma/build.py`
Результат: `docs/diploma/diploma.docx`

Форматирование делается строго по методическим указаниям ОГБПОУ
«Томский техникум информационных технологий» (файл
«Методические указания оформления текстовых документов»):
- формат А4, поля: левое 30 мм, правое 15 мм, верхнее/нижнее 20 мм;
- шрифт PT Astra Serif (fallback Times New Roman), 14 пт;
- межстрочный интервал 1.5, абзацный отступ 1.25 см;
- заголовки разделов — ПРОПИСНЫМИ, по центру, каждый раздел с новой страницы;
- заголовки подразделов — строчные с первой прописной, выравнивание слева;
- нумерация страниц — снизу по центру, арабские цифры, титульный без номера;
- списки — цифры со скобкой «1)», последний пункт заканчивается точкой.

Текст записки — в переменной DOC_CONTENT ниже. Файл писался в расчёте
на студента колледжа: простой язык, минимум теории, больше «что сделано».
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, Mm


# ---------- Низкоуровневые утилиты форматирования ----------

FONT_MAIN = 'PT Astra Serif'
FONT_FALLBACK = 'Times New Roman'


def _set_cell_borders(cell):
    """Рисует обычную чёрную границу вокруг ячейки таблицы."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:color'), '000000')
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def _apply_run_font(run, size=14, bold=False):
    run.font.name = FONT_MAIN
    run.font.size = Pt(size)
    run.bold = bold
    # Подстраховка для кириллицы (Word иногда игнорирует font.name).
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), FONT_MAIN)


def _page_number_field(paragraph):
    """Вставляет в абзац поле {PAGE} — автоматический номер страницы."""
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_end)
    _apply_run_font(run)


def _build_base_document() -> Document:
    """Создаёт пустой документ с правильными полями, шрифтом и стилями."""
    doc = Document()

    # Поля страницы.
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(30)
        section.right_margin = Mm(15)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.different_first_page_header_footer = True

    # Стиль Normal по методичке: 14 пт, 1.5 интервал, первая строка 1.25 см.
    normal = doc.styles['Normal']
    normal.font.name = FONT_MAIN
    normal.font.size = Pt(14)
    pf = normal.paragraph_format
    pf.first_line_indent = Cm(1.25)
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    return doc


def _add_paragraph(doc, text, *, bold=False, align=None, first_line_indent=None,
                   size=14, space_after=0):
    """Добавляет обычный абзац с правильным шрифтом и отступами."""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    _apply_run_font(run, size=size, bold=bold)
    return p


def _add_section_heading(doc, text, *, new_page=True):
    """Заголовок главного раздела: с новой страницы, прописные, по центру."""
    if new_page:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text.upper())
    _apply_run_font(run, size=14, bold=True)


def _add_subheading(doc, text, *, level=2):
    """Подзаголовок: слева, с первой прописной."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _apply_run_font(run, size=14, bold=True)


def _add_list(doc, items):
    """Нумерованный список в формате методички: «1)», последний — с точкой."""
    for i, item in enumerate(items, 1):
        last = i == len(items)
        suffix = '.' if last else ';'
        _add_paragraph(doc, f'{i}) {item}{suffix}', first_line_indent=Cm(1.25))


def _add_table(doc, header, rows, *, caption=None, col_widths=None):
    """Таблица с заголовком «Таблица N — Название»."""
    if caption:
        _add_paragraph(doc, caption, first_line_indent=Cm(0), space_after=4)

    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Заголовок.
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(h)
        _apply_run_font(run, size=12, bold=True)
        _set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row_idx, row in enumerate(rows, 1):
        for col_idx, val in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run(str(val))
            _apply_run_font(run, size=12)
            _set_cell_borders(cell)

    _add_paragraph(doc, '', first_line_indent=Cm(0))


def _add_code(doc, text):
    """Листинг кода: моноширинный, без абзацного отступа."""
    for line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
            rFonts.set(qn(attr), 'Courier New')


# ---------- Титульный лист ----------

def build_title_page(doc):
    """Титульный лист. Значения подставь свои (ФИО, группа, руководитель)."""
    def _center_bold(text, size=14):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(text)
        _apply_run_font(run, size=size, bold=True)

    def _center(text, size=14, bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(text)
        _apply_run_font(run, size=size, bold=bold)

    _center('МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ')
    _center('РОССИЙСКОЙ ФЕДЕРАЦИИ')
    _center('ОГБПОУ «Томский техникум информационных технологий»')
    for _ in range(6):
        _center('')
    _center_bold('ДИПЛОМНЫЙ ПРОЕКТ', size=16)
    _center('')
    _center('на тему:')
    _center_bold(
        '«Разработка автоматизированной информационной системы '
        'для работы фитнес-центра»'
    )
    for _ in range(6):
        _center('')

    # Блок «Выполнил / Руководитель» — по правому краю.
    for label, value in (
        ('Выполнил студент группы ____', '_____________________'),
        ('Руководитель', '_____________________'),
        ('Нормоконтроль', '_____________________'),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(f'{label}   {value}')
        _apply_run_font(run)

    for _ in range(6):
        _center('')
    _center('Томск 2026')
    doc.add_page_break()


# ---------- Автоматическое оглавление ----------

def build_toc(doc):
    """Добавляет поле TOC — Word сам построит оглавление после F9."""
    _add_section_heading(doc, 'Содержание', new_page=False)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    _apply_run_font(run)
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_sep)
    run._element.append(fld_end)
    _add_paragraph(
        doc,
        'Чтобы обновить оглавление в Word — выделите его и нажмите F9 '
        '(или ПКМ → «Обновить поле»).',
        first_line_indent=Cm(0), size=10,
    )


# ---------- Футер со сквозной нумерацией ----------

def build_footer(doc):
    """Ставит номер страницы снизу по центру. На первой странице не показывать."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    _page_number_field(p)
    # Первая страница (титул) — без номера.
    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    first_p = first_footer.paragraphs[0]
    first_p.text = ''


def build(output: Path) -> None:
    doc = _build_base_document()
    build_title_page(doc)
    build_toc(doc)
    # Основная часть — разбита на модули.
    from docs.diploma.content import (
        section_introduction, section_analytical, section_design,
        section_implementation, section_testing, section_conclusion,
        section_references, section_appendix_a, section_appendix_b,
        section_appendix_c,
    )
    section_introduction(doc, helpers=_helpers)
    section_analytical(doc, helpers=_helpers)
    section_design(doc, helpers=_helpers)
    section_implementation(doc, helpers=_helpers)
    section_testing(doc, helpers=_helpers)
    section_conclusion(doc, helpers=_helpers)
    section_references(doc, helpers=_helpers)
    section_appendix_a(doc, helpers=_helpers)
    section_appendix_b(doc, helpers=_helpers)
    section_appendix_c(doc, helpers=_helpers)

    build_footer(doc)
    doc.save(output)
    print(f'Готово: {output}')


# Экспортируем хелперы для модуля content.
_helpers = {
    'paragraph': _add_paragraph,
    'heading': _add_section_heading,
    'subheading': _add_subheading,
    'list': _add_list,
    'table': _add_table,
    'code': _add_code,
    'Cm': Cm, 'Pt': Pt,
    'align_center': WD_ALIGN_PARAGRAPH.CENTER,
    'align_left': WD_ALIGN_PARAGRAPH.LEFT,
}


if __name__ == '__main__':
    import sys
    # Запуск из корня проекта.
    root = Path(__file__).resolve().parent.parent.parent
    sys.path.insert(0, str(root))
    output = root / 'docs' / 'diploma' / 'diploma.docx'
    build(output)
