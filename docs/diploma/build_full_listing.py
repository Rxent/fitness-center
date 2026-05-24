"""Сборщик отдельного файла «Приложение А. Листинги исходного кода».

Создаёт самостоятельный .docx, оформленный в стиле диплома (Румянцев Р. А.,
СТАРТ КОЛЛЕДЖ, 2026): Times New Roman 11 пт, межстрочный интервал 1,0,
заголовок приложения по центру, подписи к листингам в формате
«Листинг — файл <путь>. <описание>».

Готовый файл затем можно вставить целиком вместо текущего короткого
Приложения А в основной .docx-документ диплома.

Запуск: python docs/diploma/build_full_listing.py
Результат: docs/diploma/full_listing.docx
"""
from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


FONT_NAME = 'Times New Roman'
BODY_SIZE = 14   # 14 пт для основного текста (как в дипломе)
CODE_SIZE = 11   # 11 пт для листингов (явно указано в дипломе)


# ---------------------------------------------------------------------------
# Низкоуровневые помощники форматирования
# ---------------------------------------------------------------------------

def _apply_font(run, *, size: int, bold: bool = False, italic: bool = False):
    """Назначает Times New Roman нужного размера, с подстраховкой для кириллицы."""
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), FONT_NAME)


def _build_document() -> Document:
    doc = Document()
    # А4 с теми же полями, что в дипломе: 30/15/20/20 мм.
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(30)
        section.right_margin = Mm(15)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)

    normal = doc.styles['Normal']
    normal.font.name = FONT_NAME
    normal.font.size = Pt(BODY_SIZE)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    return doc


def _add_centered_heading(doc, text: str):
    """Заголовок приложения по центру, прописными, полужирный."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    _apply_font(run, size=BODY_SIZE, bold=True)


def _add_body(doc, text: str, *, justify: bool = True):
    """Обычный абзац основного текста (14 пт, интервал 1,5, отступ первой строки)."""
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Mm(12.5)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    _apply_font(run, size=BODY_SIZE)


def _add_listing_caption(doc, rel_path: str, description: str):
    """Подпись к листингу в формате диплома: «Листинг — файл X. Описание»."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    prefix = p.add_run('Листинг — файл ')
    _apply_font(prefix, size=CODE_SIZE)
    path_run = p.add_run(rel_path)
    _apply_font(path_run, size=CODE_SIZE, italic=True)
    tail_run = p.add_run(f'. {description}.')
    _apply_font(tail_run, size=CODE_SIZE)


def _add_code(doc, source: str):
    """Тело листинга: Times New Roman 11 пт, межстрочный 1,0, без отступа."""
    for line in source.splitlines():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        # Пустые строки оставляем как пустые абзацы (чтобы Word не «съел» их).
        run = p.add_run(line if line else ' ')
        _apply_font(run, size=CODE_SIZE)


# ---------------------------------------------------------------------------
# Состав приложения: какие файлы и в каком порядке включить
# ---------------------------------------------------------------------------

FILES = [
    # Корень проекта и общие настройки.
    ('manage.py',
     'служебный скрипт управления Django-проектом'),
    ('config/settings.py',
     'настройки Django-проекта (приложения, база данных, локализация)'),
    ('config/urls.py',
     'корневой файл URL-маршрутизации'),

    # Приложение users — пользователи, аутентификация, роли.
    ('users/models.py',
     'модель пользователя с указанием роли (администратор, тренер, клиент)'),
    ('users/signals.py',
     'сигнал автоматического создания профиля клиента или тренера'),
    ('users/decorators.py',
     'декораторы для проверки роли пользователя при доступе к разделам'),
    ('users/forms.py',
     'формы регистрации клиента и редактирования профиля'),
    ('users/views.py',
     'представления для регистрации, входа, дашбордов и профиля'),
    ('users/urls.py',
     'маршруты приложения users'),
    ('users/admin.py',
     'настройка админ-панели для модели пользователя'),
    ('users/management/commands/seed_demo.py',
     'команда manage.py seed_demo для заполнения базы демо-данными'),

    # Приложение clients — клиенты фитнес-центра.
    ('clients/models.py',
     'модель клиента (профиль, привязанный к пользователю)'),
    ('clients/views.py',
     'представления списка клиентов и карточки клиента'),
    ('clients/urls.py',
     'маршруты приложения clients'),
    ('clients/admin.py',
     'настройка админ-панели для модели клиента'),

    # Приложение trainers — тренерский состав.
    ('trainers/models.py',
     'модели тренера и специализации'),
    ('trainers/views.py',
     'представления списка тренеров и карточки тренера'),
    ('trainers/urls.py',
     'маршруты приложения trainers'),
    ('trainers/admin.py',
     'настройка админ-панели для тренеров и специализаций'),

    # Приложение subscriptions — тарифы, абонементы, платежи.
    ('subscriptions/models.py',
     'модели тарифного плана, абонемента клиента и платежа'),
    ('subscriptions/views.py',
     'представления каталога тарифов и покупки абонементов'),
    ('subscriptions/urls.py',
     'маршруты приложения subscriptions'),
    ('subscriptions/admin.py',
     'настройка админ-панели для тарифов, абонементов и платежей'),

    # Приложение schedule — залы, тренировки, записи на тренировки.
    ('schedule/models.py',
     'модели залов, тренировок и записей клиентов на занятия'),
    ('schedule/views.py',
     'представления расписания, записи клиентов и отметки посещений'),
    ('schedule/urls.py',
     'маршруты приложения schedule'),
    ('schedule/admin.py',
     'настройка админ-панели для залов, тренировок и записей'),

    # Приложение reports — отчёты.
    ('reports/forms.py',
     'формы выбора периода для построения отчётов'),
    ('reports/services.py',
     'бизнес-логика расчёта отчётов (посещаемость, выручка, нагрузка тренеров)'),
    ('reports/views.py',
     'представления отчётов и страницы выбора отчёта'),
    ('reports/urls.py',
     'маршруты приложения reports'),
]


# ---------------------------------------------------------------------------
# Сборка документа
# ---------------------------------------------------------------------------

def build(output: Path) -> int:
    doc = _build_document()

    _add_centered_heading(doc, 'ПРИЛОЖЕНИЕ А. ЛИСТИНГИ ИСХОДНОГО КОДА')

    _add_body(doc,
              'В приложении приведён полный листинг исходного кода '
              'разработанной информационной системы фитнес-центра. '
              'Листинги сгруппированы по Django-приложениям проекта и '
              'упорядочены так же, как в подразделе 2.3 пояснительной '
              'записки. Файлы миграций базы данных, автоматических '
              'тестов, HTML-шаблонов и статических ресурсов в приложение '
              'не включены: миграции автоматически генерируются Django '
              'по моделям, а шаблоны и стили относятся к разметке '
              'страниц, а не к программному коду системы.')

    _add_body(doc,
              'Листинги выполнены шрифтом Times New Roman 11 пт с '
              'межстрочным интервалом 1,0. Имена файлов указаны '
              'относительно корневого каталога проекта fitness-center.')

    listings_root = Path(__file__).resolve().parent / 'source_listings'

    included = 0
    for rel_path, description in FILES:
        src_path = listings_root / rel_path
        if not src_path.exists():
            print(f'  пропущен (нет файла): {rel_path}')
            continue
        source = src_path.read_text(encoding='utf-8').rstrip('\n')
        _add_listing_caption(doc, rel_path, description)
        _add_code(doc, source)
        included += 1

    doc.save(output)
    return included


if __name__ == '__main__':
    output = Path(__file__).resolve().parent / 'full_listing.docx'
    count = build(output)
    print(f'Готово: {output}')
    print(f'Включено листингов: {count}')
